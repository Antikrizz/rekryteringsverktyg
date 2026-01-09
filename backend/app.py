from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import os
import json
import sqlite3
from datetime import datetime
import anthropic
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import io
import base64
import tempfile

# Ladda miljövariabler
load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))

app = Flask(__name__)
CORS(app)

# API-klienter
anthropic_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Databas-setup
DB_PATH = os.path.join(os.path.dirname(__file__), 'rekrytering.db')

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute('''CREATE TABLE IF NOT EXISTS roles (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        description TEXT,
        questions TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS candidates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        role_id INTEGER,
        cv_text TEXT,
        personal_questions TEXT,
        all_questions TEXT,
        transcript TEXT,
        analysis TEXT,
        total_score INTEGER,
        interview_date TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (role_id) REFERENCES roles (id)
    )''')

    conn.commit()
    conn.close()

init_db()

# Hjälpfunktioner
def get_db():
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute('PRAGMA journal_mode=WAL')
    return conn

def extract_text_from_pdf(file):
    """Extrahera text från PDF-fil"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Fel vid läsning av PDF: {str(e)}"

def extract_text_from_docx(file):
    """Extrahera text från Word-fil"""
    try:
        doc = Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return f"Fel vid läsning av Word-fil: {str(e)}"

# API Routes

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({"status": "ok", "message": "Backend körs!"})

# === ROLLER ===

@app.route('/api/roles', methods=['GET'])
def get_roles():
    conn = get_db()
    roles = conn.execute('SELECT * FROM roles ORDER BY created_at DESC').fetchall()
    conn.close()
    return jsonify([dict(row) for row in roles])

@app.route('/api/roles', methods=['POST'])
def create_role():
    data = request.json
    name = data.get('name')
    description = data.get('description', '')

    if not name:
        return jsonify({"error": "Rollnamn krävs"}), 400

    # Generera frågor med Claude
    questions = generate_role_questions(name, description)

    conn = get_db()
    cursor = conn.execute(
        'INSERT INTO roles (name, description, questions) VALUES (?, ?, ?)',
        (name, description, json.dumps(questions, ensure_ascii=False))
    )
    role_id = cursor.lastrowid
    conn.commit()
    conn.close()

    return jsonify({
        "id": role_id,
        "name": name,
        "description": description,
        "questions": questions
    })

@app.route('/api/roles/<int:role_id>', methods=['GET'])
def get_role(role_id):
    conn = get_db()
    role = conn.execute('SELECT * FROM roles WHERE id = ?', (role_id,)).fetchone()
    conn.close()

    if not role:
        return jsonify({"error": "Roll hittades inte"}), 404

    role_dict = dict(role)
    role_dict['questions'] = json.loads(role_dict['questions']) if role_dict['questions'] else []
    return jsonify(role_dict)

@app.route('/api/roles/<int:role_id>', methods=['PUT'])
def update_role(role_id):
    data = request.json
    questions = data.get('questions', [])

    conn = get_db()
    conn.execute(
        'UPDATE roles SET questions = ? WHERE id = ?',
        (json.dumps(questions, ensure_ascii=False), role_id)
    )
    conn.commit()
    conn.close()

    return jsonify({"success": True})

@app.route('/api/roles/<int:role_id>', methods=['DELETE'])
def delete_role(role_id):
    conn = get_db()
    conn.execute('DELETE FROM roles WHERE id = ?', (role_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# === KANDIDATER ===

@app.route('/api/candidates', methods=['GET'])
def get_candidates():
    conn = get_db()
    candidates = conn.execute('''
        SELECT c.*, r.name as role_name
        FROM candidates c
        LEFT JOIN roles r ON c.role_id = r.id
        ORDER BY c.total_score DESC, c.created_at DESC
    ''').fetchall()
    conn.close()
    return jsonify([dict(row) for row in candidates])

@app.route('/api/candidates/<int:candidate_id>', methods=['GET'])
def get_candidate(candidate_id):
    conn = get_db()
    candidate = conn.execute('''
        SELECT c.*, r.name as role_name, r.description as role_description
        FROM candidates c
        LEFT JOIN roles r ON c.role_id = r.id
        WHERE c.id = ?
    ''', (candidate_id,)).fetchone()
    conn.close()

    if not candidate:
        return jsonify({"error": "Kandidat hittades inte"}), 404

    candidate_dict = dict(candidate)
    if candidate_dict['analysis']:
        candidate_dict['analysis'] = json.loads(candidate_dict['analysis'])
    if candidate_dict['all_questions']:
        candidate_dict['all_questions'] = json.loads(candidate_dict['all_questions'])
    if candidate_dict['personal_questions']:
        candidate_dict['personal_questions'] = json.loads(candidate_dict['personal_questions'])

    return jsonify(candidate_dict)

@app.route('/api/candidates/<int:candidate_id>', methods=['DELETE'])
def delete_candidate(candidate_id):
    conn = get_db()
    conn.execute('DELETE FROM candidates WHERE id = ?', (candidate_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# === CV UPLOAD & PERSONAL QUESTIONS ===

@app.route('/api/upload-cv', methods=['POST'])
def upload_cv():
    if 'file' not in request.files:
        # Kolla om det är text direkt
        data = request.json
        if data and data.get('cv_text'):
            cv_text = data.get('cv_text')
        else:
            return jsonify({"error": "Ingen fil eller text skickad"}), 400
    else:
        file = request.files['file']
        filename = file.filename.lower()

        if filename.endswith('.pdf'):
            cv_text = extract_text_from_pdf(file)
        elif filename.endswith('.docx'):
            cv_text = extract_text_from_docx(file)
        elif filename.endswith('.txt'):
            cv_text = file.read().decode('utf-8')
        else:
            return jsonify({"error": "Filformat stöds inte. Använd PDF, Word eller TXT."}), 400

    return jsonify({"cv_text": cv_text})

@app.route('/api/generate-personal-questions', methods=['POST'])
def generate_personal_questions():
    data = request.json
    cv_text = data.get('cv_text', '')
    role_description = data.get('role_description', '')
    role_name = data.get('role_name', '')

    if not cv_text:
        return jsonify({"error": "CV-text krävs"}), 400

    questions = generate_cv_questions(cv_text, role_name, role_description)
    return jsonify({"questions": questions})

# === INTERVJU & ANALYS ===

@app.route('/api/prepare-candidate', methods=['POST'])
def prepare_candidate():
    data = request.json
    role_id = data.get('role_id')
    cv_text = data.get('cv_text', '')
    personal_questions = data.get('personal_questions', [])

    conn = get_db()
    role = conn.execute('SELECT * FROM roles WHERE id = ?', (role_id,)).fetchone()

    if not role:
        conn.close()
        return jsonify({"error": "Roll hittades inte"}), 404

    role_questions = json.loads(role['questions']) if role['questions'] else []
    all_questions = role_questions + personal_questions

    cursor = conn.execute(
        '''INSERT INTO candidates (role_id, cv_text, personal_questions, all_questions)
           VALUES (?, ?, ?, ?)''',
        (role_id, cv_text, json.dumps(personal_questions, ensure_ascii=False),
         json.dumps(all_questions, ensure_ascii=False))
    )
    candidate_id = cursor.lastrowid
    conn.commit()
    conn.close()

    return jsonify({
        "candidate_id": candidate_id,
        "all_questions": all_questions
    })

@app.route('/api/transcribe', methods=['POST'])
def transcribe_audio():
    if 'file' not in request.files:
        return jsonify({"error": "Ingen ljudfil skickad"}), 400

    file = request.files['file']

    try:
        # Spara temporärt
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        # Transkribera med Whisper
        with open(tmp_path, 'rb') as audio_file:
            transcript = openai_client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language="sv"
            )

        # Ta bort temp-fil
        os.unlink(tmp_path)

        return jsonify({"transcript": transcript.text})

    except Exception as e:
        return jsonify({"error": f"Transkribering misslyckades: {str(e)}"}), 500

@app.route('/api/analyze-interview', methods=['POST'])
def analyze_interview():
    data = request.json
    candidate_id = data.get('candidate_id')
    candidate_name = data.get('candidate_name')
    transcript = data.get('transcript', '')

    conn = get_db()
    candidate = conn.execute(
        'SELECT c.*, r.name as role_name FROM candidates c JOIN roles r ON c.role_id = r.id WHERE c.id = ?',
        (candidate_id,)
    ).fetchone()

    if not candidate:
        conn.close()
        return jsonify({"error": "Kandidat hittades inte"}), 404

    all_questions = json.loads(candidate['all_questions'])

    # Analysera med Claude
    analysis = analyze_with_claude(all_questions, transcript, candidate['role_name'])

    # Beräkna totalpoäng
    total_score = sum(q.get('score', 0) for q in analysis.get('questions', []))

    # Uppdatera kandidat
    conn.execute(
        '''UPDATE candidates SET
           name = ?, transcript = ?, analysis = ?, total_score = ?, interview_date = ?
           WHERE id = ?''',
        (candidate_name, transcript, json.dumps(analysis, ensure_ascii=False),
         total_score, datetime.now().isoformat(), candidate_id)
    )
    conn.commit()
    conn.close()

    return jsonify({
        "analysis": analysis,
        "total_score": total_score
    })

# === RAPPORT ===

@app.route('/api/report/<int:candidate_id>', methods=['GET'])
def generate_report(candidate_id):
    conn = get_db()
    candidate = conn.execute('''
        SELECT c.*, r.name as role_name
        FROM candidates c
        LEFT JOIN roles r ON c.role_id = r.id
        WHERE c.id = ?
    ''', (candidate_id,)).fetchone()
    conn.close()

    if not candidate:
        return jsonify({"error": "Kandidat hittades inte"}), 404

    candidate_dict = dict(candidate)
    analysis = json.loads(candidate_dict['analysis']) if candidate_dict['analysis'] else {}

    # Skapa Word-dokument
    doc = Document()

    # Titel
    title = doc.add_heading('Intervjurapport', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Grundinformation
    doc.add_heading('Grundinformation', level=1)
    doc.add_paragraph(f"Roll: {candidate_dict.get('role_name', 'Ej angiven')}")
    doc.add_paragraph(f"Kandidat: {candidate_dict.get('name', 'Ej angiven')}")
    doc.add_paragraph(f"Datum: {candidate_dict.get('interview_date', '')[:10] if candidate_dict.get('interview_date') else 'Ej angivet'}")
    doc.add_paragraph(f"Totalpoäng: {candidate_dict.get('total_score', 0)}/50")

    # Övergripande bedömning
    doc.add_heading('Övergripande bedömning', level=1)
    doc.add_paragraph(analysis.get('overall_assessment', 'Ingen bedömning tillgänglig.'))

    # Frågor och bedömning
    doc.add_heading('Frågor och bedömning', level=1)

    for i, q in enumerate(analysis.get('questions', []), 1):
        doc.add_heading(f"Fråga {i}", level=2)
        doc.add_paragraph(q.get('question', ''))
        doc.add_paragraph(f"Poäng: {q.get('score', 0)}/5")
        doc.add_paragraph(f"Sammanfattning: {q.get('summary', '')}")
        doc.add_paragraph(f"Bedömning: {q.get('assessment', '')}")
        if q.get('quote'):
            doc.add_paragraph(f'Citat: "{q.get("quote", "")}"')
        doc.add_paragraph()

    # Sammanfattad transkription
    doc.add_heading('Sammanfattad transkription', level=1)
    doc.add_paragraph(analysis.get('summarized_transcript', 'Ingen transkription tillgänglig.'))

    # Spara till bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    filename = f"intervjurapport_{candidate_dict.get('name', 'kandidat').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"

    return send_file(
        doc_bytes,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

# === AI-FUNKTIONER ===

def generate_role_questions(role_name, role_description):
    """Generera 6 intervjufrågor baserat på roll"""

    prompt = f"""Du är en expert på rekrytering. Generera 6 intervjufrågor för rollen "{role_name}".

Rollbeskrivning:
{role_description if role_description else "Ingen specifik rollbeskrivning angiven."}

Frågorna ska täcka dessa kategorier (anpassa efter rollen):
1. Teknisk kompetens och erfarenhet
2. Ledarskap och förändring
3. Teambuilding och samarbete
4. Affärsmässighet och kundkontakt
5. Innovation, AI och digitalisering
6. Hållbarhet och framtidsperspektiv

Svara ENDAST med en JSON-array med 6 objekt. Varje objekt ska ha:
- "category": kategorinamn
- "question": själva frågan

Exempel på format:
[
  {{"category": "Teknisk kompetens", "question": "Berätta om din tekniska bakgrund..."}},
  ...
]

Svara ENDAST med JSON-arrayen, inget annat."""

    try:
        response = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text.strip()
        # Ta bort eventuella markdown-kodblock
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]

        questions = json.loads(response_text)
        return questions

    except Exception as e:
        print(f"Fel vid generering av frågor: {e}")
        # Returnera standardfrågor om något går fel
        return [
            {"category": "Teknisk kompetens", "question": "Berätta om din tekniska kompetens och dina erfarenheter från tidigare arbeten."},
            {"category": "Ledarskap", "question": "Hur ser du på ledarskap och hur skulle du beskriva din ledarstil?"},
            {"category": "Teambuilding", "question": "Hur skulle du gå tillväga för att skapa en stark teamkultur?"},
            {"category": "Affärsmässighet", "question": "Hur ser du på affärsutveckling och kundrelationer?"},
            {"category": "Innovation", "question": "Hur ser du på möjligheterna med AI och digitalisering inom ditt område?"},
            {"category": "Hållbarhet", "question": "Hur ser du på hållbarhet och miljöpåverkan i arbetet?"}
        ]

def generate_cv_questions(cv_text, role_name, role_description):
    """Generera 4 personliga frågor baserat på CV"""

    prompt = f"""Du är en expert på rekrytering. Analysera detta CV och generera 4 personliga intervjufrågor.

Roll: {role_name}
Rollbeskrivning: {role_description if role_description else "Ej angiven"}

CV:
{cv_text[:5000]}

Generera 4 specifika frågor baserade på:
- Kandidatens tidigare erfarenheter
- Luckor eller intressanta punkter i CV:t
- Hur kandidatens bakgrund matchar rollen
- Specifika projekt eller prestationer att fördjupa

Svara ENDAST med en JSON-array med 4 objekt:
[
  {{"category": "Personlig", "question": "Din fråga här..."}},
  ...
]

Svara ENDAST med JSON-arrayen, inget annat."""

    try:
        response = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text.strip()
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]

        questions = json.loads(response_text)
        return questions

    except Exception as e:
        print(f"Fel vid generering av CV-frågor: {e}")
        return [
            {"category": "Personlig", "question": "Berätta mer om din senaste arbetsplats och vad du lärde dig där."},
            {"category": "Personlig", "question": "Vad motiverade dig att söka denna tjänst?"},
            {"category": "Personlig", "question": "Vilken är din största professionella prestation?"},
            {"category": "Personlig", "question": "Var ser du dig själv om 5 år?"}
        ]

def analyze_with_claude(questions, transcript, role_name):
    """Analysera intervju med Claude"""

    questions_text = "\n".join([f"{i+1}. [{q.get('category', '')}] {q.get('question', '')}"
                                 for i, q in enumerate(questions)])

    prompt = f"""Du är en expert på rekrytering och ska analysera en intervju för rollen "{role_name}".

INTERVJUFRÅGOR:
{questions_text}

TRANSKRIPTION AV INTERVJUN:
{transcript}

UPPGIFT:
1. Matcha kandidatens svar till rätt frågor (svaren kan komma i annan ordning)
2. Bedöm varje svar på skala 1-5:
   - 5: Exceptionellt - djup förståelse, konkreta exempel, strategiskt tänkande
   - 4: Starkt - tydlig kompetens, relevanta exempel
   - 3: Acceptabelt - grundläggande förståelse, saknar djup
   - 2: Svagt - vag eller bristfällig
   - 1: Mycket svagt - ingen relevant förståelse

Svara med JSON i exakt detta format:
{{
  "overall_assessment": "3-4 meningars övergripande bedömning av kandidaten",
  "summarized_transcript": "Sammanfattning av hela intervjun (max 200 ord)",
  "questions": [
    {{
      "question": "Frågetexten",
      "score": 4,
      "summary": "Kort sammanfattning av svaret",
      "assessment": "Motivering till poängen",
      "quote": "Ett kort citat från kandidaten (max 20 ord)"
    }}
  ]
}}

VIKTIGT:
- Inkludera alla {len(questions)} frågor i svaret
- Svara ENDAST med JSON, inget annat"""

    try:
        response = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text.strip()
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]

        analysis = json.loads(response_text)
        return analysis

    except Exception as e:
        print(f"Fel vid analys: {e}")
        return {
            "overall_assessment": "Analysen kunde inte genomföras på grund av ett tekniskt fel.",
            "summarized_transcript": transcript[:500] + "..." if len(transcript) > 500 else transcript,
            "questions": [{"question": q.get('question', ''), "score": 0, "summary": "Kunde inte analyseras",
                          "assessment": "Tekniskt fel", "quote": ""} for q in questions]
        }

if __name__ == '__main__':
    print("Startar backend på http://localhost:5000")
    app.run(debug=True, port=5000)
